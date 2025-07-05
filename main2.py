#Import dan Setup Awal
import tkinter as tk
from tkinter import messagebox
import speech_recognition as sr
from docx import Document
from datetime import datetime
import threading

# Membuat file Word baru
doc = Document()
doc.add_heading('Hasil Pengenalan Suara', 0)

# Inisialisasi recognizer
recognizer = sr.Recognizer()

# Variabel global untuk kontrol perekaman
is_recording = False

# Fungsi untuk menambahkan teks ke dokumen Word
def add_to_word(text):
    doc.add_paragraph(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {text}")
    doc.save("hasil_suara.docx")
    print(f"Teks ditambahkan ke Word: {text}")
    text_display.insert(tk.END, f"{text}\n")
    text_display.see(tk.END)  # Scroll ke bawah

# Fungsi untuk merekam suara terus-menerus
def start_recording():
    global is_recording
    is_recording = True
    status_label.config(text="Mendengarkan...", fg="green")
    
    def record():
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source)  # Menyesuaikan dengan kebisingan latar
            while is_recording:
                try:
                    audio = recognizer.listen(source)  # Mendengarkan suara
                    text = recognizer.recognize_google(audio, language='id-ID')  # Bahasa Indonesia
                    add_to_word(text)
                except sr.UnknownValueError:
                    pass  # Mengabaikan error jika tidak bisa memahami suara
                except sr.RequestError as e:
                    status_label.config(text=f"Permintaan gagal; {e}", fg="red")
                    break
    
    # Menjalankan perekaman dalam thread terpisah agar UI tetap responsif
    thread = threading.Thread(target=record)
    thread.daemon = True  # Thread akan berhenti saat program berhenti
    thread.start()

# Fungsi untuk menghentikan perekaman
def stop_recording():
    global is_recording
    is_recording = False
    status_label.config(text="Perekaman dihentikan.", fg="red")

# Fungsi untuk keluar dari aplikasi
def exit_application():
    if messagebox.askyesno("Konfirmasi", "Apakah Anda yakin ingin keluar?"):
        root.quit()

# Membuat aplikasi Tkinter
root = tk.Tk()
root.title("Aplikasi Pengenalan Suara")

# Membuat Label status
status_label = tk.Label(root, text="Tekan 'Mulai' untuk mulai mendengarkan", font=("Arial", 14))
status_label.pack(pady=10)

# Membuat Textbox untuk menampilkan hasil teks
text_display = tk.Text(root, height=10, width=50)
text_display.pack(pady=10)

# Tombol untuk mulai merekam suara
start_button = tk.Button(root, text="Mulai Rekam", font=("Arial", 14), command=start_recording)
start_button.pack(pady=10)

# Tombol untuk menghentikan perekaman
stop_button = tk.Button(root, text="Stop Rekam", font=("Arial", 14), command=stop_recording)
stop_button.pack(pady=10)

# Tombol untuk keluar dari aplikasi
exit_button = tk.Button(root, text="Keluar", font=("Arial", 14), command=exit_application)
exit_button.pack(pady=10)

# Menjalankan aplikasi Tkinter
root.mainloop()