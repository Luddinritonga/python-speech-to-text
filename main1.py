import speech_recognition as sr
from docx import Document

# Membuat file Word baru
doc = Document()
doc.add_heading('Hasil Pengenalan Suara dari tugas saya', 0)

# Inisialisasi recognizer
recognizer = sr.Recognizer()

# Fungsi untuk menambahkan teks ke dokumen Word
def add_to_word(text):
    doc.add_paragraph(text)
    doc.save("rekaman.docx")
    print(f"Teks ditambahkan ke Word: {text}")

# Menggunakan mikrofon sebagai sumber suara
with sr.Microphone() as source:
    print("Mendengarkan suara... (tekan Ctrl+C untuk berhenti)")
    
    while True:
        try:
            recognizer.adjust_for_ambient_noise(source)  # Menyesuaikan dengan kebisingan latar
            audio = recognizer.listen(source)  # Mendengarkan suara
            print("Mengubah suara menjadi teks...")
            
            # Mengubah suara menjadi teks dan menambahkannya ke Word
            text = recognizer.recognize_google(audio, language='id-ID')  # Bahasa Indonesia
            add_to_word(text)
        
        except sr.UnknownValueError:
            print("Google Speech Recognition tidak bisa memahami audio")
        except sr.RequestError as e:
            print(f"Permintaan ke Google Speech Recognition gagal; {e}")
        except KeyboardInterrupt:
            print("Program dihentikan.")
            break